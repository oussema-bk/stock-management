#!/usr/bin/env python3
"""
Generate comprehensive documentation for the Electrical Parts Agency Django application
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

def create_documentation():
    """Create comprehensive documentation for the application"""
    
    # Create a new Document
    doc = Document()
    
    # Title
    title = doc.add_heading('Electrical Parts Agency - Stock Management System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_heading('Complete Django Application Documentation', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc = [
        "1. System Overview",
        "2. Technical Architecture",
        "3. Features and Functionalities",
        "4. Database Schema",
        "5. API Endpoints",
        "6. User Guide",
        "7. Installation and Setup",
        "8. Testing and Validation",
        "9. Deployment Guidelines",
        "10. Troubleshooting"
    ]
    
    for item in toc:
        doc.add_paragraph(item, style='List Number')
    
    # Page break
    doc.add_page_break()
    
    # 1. System Overview
    doc.add_heading('1. System Overview', level=1)
    
    doc.add_paragraph(
        "The Electrical Parts Agency Stock Management System is a comprehensive Django-based web application "
        "designed to manage inventory, sales, and stock movements for an electrical parts distribution business. "
        "The system provides a complete solution for tracking products, managing stock levels, processing sales, "
        "and generating reports."
    )
    
    doc.add_paragraph("Key Features:")
    features = [
        "Product catalog management with categories and SKU tracking",
        "Real-time stock level monitoring with low-stock alerts",
        "Comprehensive sales tracking and customer management",
        "Stock movement tracking (in/out/adjustments)",
        "Responsive web interface with Bootstrap 5",
        "Admin panel for system administration",
        "Data visualization with charts and graphs",
        "Automated fake data generation for testing"
    ]
    
    for feature in features:
        doc.add_paragraph(f"• {feature}", style='List Bullet')
    
    # 2. Technical Architecture
    doc.add_heading('2. Technical Architecture', level=1)
    
    doc.add_paragraph("The application is built using the following technology stack:")
    
    tech_stack = [
        ("Backend Framework", "Django 4.2.7"),
        ("Database", "SQLite (development)"),
        ("Frontend", "Bootstrap 5, HTML5, CSS3, JavaScript"),
        ("Charts", "Chart.js"),
        ("Authentication", "Django built-in authentication"),
        ("Image Handling", "Pillow"),
        ("Fake Data", "Faker library"),
        ("Documentation", "python-docx")
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Technology'
    
    for component, technology in tech_stack:
        row_cells = table.add_row().cells
        row_cells[0].text = component
        row_cells[1].text = technology
    
    # Application Structure
    doc.add_heading('2.1 Application Structure', level=2)
    
    doc.add_paragraph("The Django project consists of the following apps:")
    
    apps = [
        ("products_app", "Product catalog management, categories, and product details"),
        ("stock_app", "Stock level tracking, movements, and alerts"),
        ("sales_app", "Sales processing, customer management, and analytics"),
        ("dashboard_app", "Main dashboard with overview statistics"),
        ("accounts_app", "User authentication and profile management")
    ]
    
    for app, description in apps:
        doc.add_paragraph(f"• {app}: {description}", style='List Bullet')
    
    # 3. Features and Functionalities
    doc.add_heading('3. Features and Functionalities', level=1)
    
    # 3.1 Product Management
    doc.add_heading('3.1 Product Management', level=2)
    
    doc.add_paragraph("The product management module provides comprehensive functionality for managing electrical parts inventory:")
    
    product_features = [
        "Create, read, update, and delete products",
        "Product categorization with custom categories",
        "SKU (Stock Keeping Unit) management",
        "Price and cost price tracking with profit margin calculation",
        "Product image upload and management",
        "Product status management (active/inactive)",
        "Search and filter functionality",
        "Bulk operations support"
    ]
    
    for feature in product_features:
        doc.add_paragraph(f"• {feature}", style='List Bullet')
    
    # 3.2 Stock Management
    doc.add_heading('3.2 Stock Management', level=2)
    
    doc.add_paragraph("The stock management module tracks inventory levels and movements:")
    
    stock_features = [
        "Real-time stock level monitoring",
        "Stock movement tracking (in, out, adjustments, transfers)",
        "Low stock and out-of-stock alerts",
        "Minimum and maximum stock level configuration",
        "Stock movement history and audit trail",
        "Automated stock updates based on sales",
        "Stock level reports and analytics"
    ]
    
    for feature in stock_features:
        doc.add_paragraph(f"• {feature}", style='List Bullet')
    
    # 3.3 Sales Management
    doc.add_heading('3.3 Sales Management', level=2)
    
    doc.add_paragraph("The sales management module handles customer transactions and sales analytics:")
    
    sales_features = [
        "Customer management and contact information",
        "Sales order creation and processing",
        "Multi-item sales support",
        "Sales status tracking (pending, completed, cancelled)",
        "Sales analytics and reporting",
        "Top-selling products tracking",
        "Revenue and sales performance metrics",
        "Customer sales history"
    ]
    
    for feature in sales_features:
        doc.add_paragraph(f"• {feature}", style='List Bullet')
    
    # 3.4 Dashboard and Analytics
    doc.add_heading('3.4 Dashboard and Analytics', level=2)
    
    doc.add_paragraph("The dashboard provides a comprehensive overview of business operations:")
    
    dashboard_features = [
        "Key performance indicators (KPIs)",
        "Stock status overview",
        "Sales performance metrics",
        "Low stock alerts",
        "Recent activity feed",
        "Quick action buttons",
        "Revenue tracking",
        "Top-selling products display"
    ]
    
    for feature in dashboard_features:
        doc.add_paragraph(f"• {feature}", style='List Bullet')
    
    # 4. Database Schema
    doc.add_heading('4. Database Schema', level=1)
    
    doc.add_paragraph("The application uses the following main database models:")
    
    # Products App Models
    doc.add_heading('4.1 Products App Models', level=2)
    
    models_info = [
        ("Category", "Stores product categories with name and description"),
        ("Product", "Main product model with SKU, pricing, and category relationship")
    ]
    
    for model, description in models_info:
        doc.add_paragraph(f"• {model}: {description}", style='List Bullet')
    
    # Stock App Models
    doc.add_heading('4.2 Stock App Models', level=2)
    
    stock_models = [
        ("StockLevel", "Tracks current stock levels, minimum/maximum thresholds"),
        ("StockMovement", "Records all stock movements with type, quantity, and reference")
    ]
    
    for model, description in stock_models:
        doc.add_paragraph(f"• {model}: {description}", style='List Bullet')
    
    # Sales App Models
    doc.add_heading('4.3 Sales App Models', level=2)
    
    sales_models = [
        ("Customer", "Customer information and contact details"),
        ("Sale", "Sales transactions with status and total amount"),
        ("SaleItem", "Individual items within a sale with quantity and pricing")
    ]
    
    for model, description in sales_models:
        doc.add_paragraph(f"• {model}: {description}", style='List Bullet')
    
    # 5. API Endpoints
    doc.add_heading('5. API Endpoints', level=1)
    
    doc.add_paragraph("The application provides the following main URL patterns:")
    
    # Products URLs
    doc.add_heading('5.1 Products URLs', level=2)
    
    product_urls = [
        ("/products/", "GET", "Product list view with search and filtering"),
        ("/products/create/", "GET, POST", "Create new product"),
        ("/products/<id>/", "GET", "Product detail view"),
        ("/products/<id>/update/", "GET, POST", "Update product"),
        ("/products/<id>/delete/", "GET, POST", "Delete product"),
        ("/products/categories/", "GET", "Category list view"),
        ("/products/categories/create/", "GET, POST", "Create new category")
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'URL'
    hdr_cells[1].text = 'Method'
    hdr_cells[2].text = 'Description'
    
    for url, method, description in product_urls:
        row_cells = table.add_row().cells
        row_cells[0].text = url
        row_cells[1].text = method
        row_cells[2].text = description
    
    # Stock URLs
    doc.add_heading('5.2 Stock URLs', level=2)
    
    stock_urls = [
        ("/stock/", "GET", "Stock dashboard with overview"),
        ("/stock/list/", "GET", "Stock level list with filtering"),
        ("/stock/movements/", "GET", "Stock movements history"),
        ("/stock/movements/add/", "GET, POST", "Add new stock movement"),
        ("/stock/level/<id>/update/", "GET, POST", "Update stock level settings"),
        ("/stock/alerts/", "GET", "Low stock alerts view"),
        ("/stock/api/", "GET", "Stock data API endpoint")
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'URL'
    hdr_cells[1].text = 'Method'
    hdr_cells[2].text = 'Description'
    
    for url, method, description in stock_urls:
        row_cells = table.add_row().cells
        row_cells[0].text = url
        row_cells[1].text = method
        row_cells[2].text = description
    
    # Sales URLs
    doc.add_heading('5.3 Sales URLs', level=2)
    
    sales_urls = [
        ("/sales/", "GET", "Sales dashboard with analytics"),
        ("/sales/sales/", "GET", "Sales list with filtering"),
        ("/sales/sales/create/", "GET, POST", "Create new sale"),
        ("/sales/sales/<id>/", "GET", "Sale detail view"),
        ("/sales/sales/<id>/add-item/", "GET, POST", "Add item to sale"),
        ("/sales/customers/", "GET", "Customer list"),
        ("/sales/customers/create/", "GET, POST", "Create new customer"),
        ("/sales/analytics/", "GET", "Sales analytics and charts"),
        ("/sales/api/", "GET", "Sales data API endpoint")
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'URL'
    hdr_cells[1].text = 'Method'
    hdr_cells[2].text = 'Description'
    
    for url, method, description in sales_urls:
        row_cells = table.add_row().cells
        row_cells[0].text = url
        row_cells[1].text = method
        row_cells[2].text = description
    
    # 6. User Guide
    doc.add_heading('6. User Guide', level=1)
    
    # 6.1 Admin User Guide
    doc.add_heading('6.1 Admin User Guide', level=2)
    
    doc.add_paragraph("As an admin user, you have access to all system features:")
    
    admin_guide = [
        "Access the admin panel at /admin/ to manage all data",
        "Use the main dashboard at /dashboard/ for business overview",
        "Navigate through products, stock, and sales modules using the sidebar",
        "Create and manage product categories before adding products",
        "Set up stock levels and minimum/maximum thresholds",
        "Process sales and track customer information",
        "Monitor low stock alerts and take appropriate actions",
        "Generate reports and analyze business performance"
    ]
    
    for item in admin_guide:
        doc.add_paragraph(f"• {item}", style='List Bullet')
    
    # 6.2 Staff User Guide
    doc.add_heading('6.2 Staff User Guide', level=2)
    
    doc.add_paragraph("Staff users can perform day-to-day operations:")
    
    staff_guide = [
        "View and search products in the catalog",
        "Process stock movements (in/out/adjustments)",
        "Create and manage sales transactions",
        "Add new customers and update their information",
        "Monitor stock levels and respond to alerts",
        "View sales reports and analytics",
        "Update product information as needed"
    ]
    
    for item in staff_guide:
        doc.add_paragraph(f"• {item}", style='List Bullet')
    
    # 7. Installation and Setup
    doc.add_heading('7. Installation and Setup', level=1)
    
    doc.add_paragraph("Follow these steps to set up the application:")
    
    setup_steps = [
        "Ensure Python 3.8+ is installed on your system",
        "Install required packages: pip install django pillow faker python-docx django-plotly-dash black",
        "Run database migrations: python manage.py migrate",
        "Create a superuser: python manage.py createsuperuser",
        "Generate fake data: python manage.py generate_fake_data",
        "Collect static files: python manage.py collectstatic",
        "Start the development server: python manage.py runserver",
        "Access the application at http://localhost:8000/"
    ]
    
    for i, step in enumerate(setup_steps, 1):
        doc.add_paragraph(f"{i}. {step}", style='List Number')
    
    # Automated Setup
    doc.add_heading('7.1 Automated Setup', level=2)
    
    doc.add_paragraph(
        "For convenience, use the provided setup script:"
    )
    
    doc.add_paragraph("chmod +x setup.sh")
    doc.add_paragraph("./setup.sh")
    
    doc.add_paragraph(
        "This script will automatically install dependencies, run migrations, create a superuser, "
        "generate fake data, and run tests."
    )
    
    # 8. Testing and Validation
    doc.add_heading('8. Testing and Validation', level=1)
    
    doc.add_paragraph("The application includes comprehensive testing:")
    
    testing_info = [
        "Automated test suite (test_app.py) validates all core functionality",
        "Database model testing ensures data integrity",
        "View testing verifies all pages are accessible",
        "Admin interface testing confirms administrative access",
        "Data relationship testing validates model connections",
        "Fake data generation provides realistic test scenarios"
    ]
    
    for item in testing_info:
        doc.add_paragraph(f"• {item}", style='List Bullet')
    
    doc.add_paragraph("Run tests using: python3 test_app.py")
    
    # 9. Deployment Guidelines
    doc.add_heading('9. Deployment Guidelines', level=1)
    
    doc.add_paragraph("For production deployment:")
    
    deployment_steps = [
        "Set DEBUG = False in settings.py",
        "Configure a production database (PostgreSQL recommended)",
        "Set up proper static file serving",
        "Configure media file storage",
        "Set up SSL/HTTPS",
        "Configure proper ALLOWED_HOSTS",
        "Set up logging and monitoring",
        "Use a production WSGI server (Gunicorn recommended)",
        "Set up reverse proxy (Nginx recommended)"
    ]
    
    for i, step in enumerate(deployment_steps, 1):
        doc.add_paragraph(f"{i}. {step}", style='List Number')
    
    # 10. Troubleshooting
    doc.add_heading('10. Troubleshooting', level=1)
    
    doc.add_paragraph("Common issues and solutions:")
    
    troubleshooting = [
        ("Database errors", "Run python manage.py migrate to apply migrations"),
        ("Static files not loading", "Run python manage.py collectstatic"),
        ("Permission errors", "Check file permissions and ownership"),
        ("Import errors", "Ensure all dependencies are installed"),
        ("Template not found", "Check TEMPLATES setting in settings.py"),
        ("Media files not serving", "Configure MEDIA_URL and MEDIA_ROOT"),
        ("CSRF errors", "Ensure CSRF middleware is enabled"),
        ("Login issues", "Verify user exists and password is correct")
    ]
    
    for issue, solution in troubleshooting:
        doc.add_paragraph(f"• {issue}: {solution}", style='List Bullet')
    
    # Conclusion
    doc.add_heading('Conclusion', level=1)
    
    doc.add_paragraph(
        "The Electrical Parts Agency Stock Management System provides a comprehensive solution for managing "
        "electrical parts inventory, sales, and stock movements. The system is built with Django best practices, "
        "includes responsive design, and provides extensive functionality for both administrative and operational users."
    )
    
    doc.add_paragraph(
        "The application is ready for immediate use and can be easily extended with additional features as needed. "
        "All code is well-documented and follows Django conventions for maintainability and scalability."
    )
    
    # Save the document
    doc.save('Electrical_Parts_Agency_Documentation.docx')
    print("Documentation generated successfully: Electrical_Parts_Agency_Documentation.docx")

if __name__ == "__main__":
    create_documentation()
